import pandas as pd
from docx import Document

def load_data(excel_path: str) -> pd.DataFrame:
    """Load the ‘All_Papers’ sheet into a DataFrame."""
    return pd.read_excel(excel_path, sheet_name='All_Papers')

def add_general_section(doc: Document, df: pd.DataFrame):
    """Add the ‘General Bug Detection’ section to the Word doc."""
    doc.add_heading('1A. General Bug Detection Papers', level=2)
    general = df[df['primary_bug_type'] == 'General']
    for _, row in general.iterrows():
        dataset = row.get('dataset_used') or row.get('dataset_or_tool_link') or "Unknown"
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{row['paper_title']}: Dataset = {dataset}")

def add_specific_section(doc: Document, df: pd.DataFrame):
    """Add the ‘Specific Bug Categories’ section, grouped by bug type."""
    doc.add_heading('1B. Specific Bug Categories', level=2)
    specific = df[df['primary_bug_type'] != 'General']
    for bug_type, group in specific.groupby('primary_bug_type'):
        doc.add_heading(bug_type, level=3)
        for _, row in group.iterrows():
            dataset = row.get('dataset_used') or row.get('dataset_or_tool_link') or "Unknown"
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{row['paper_title']}: Dataset = {dataset}")

def add_integration_section(doc: Document, df: pd.DataFrame):
    """Add the ‘LLM Role and Integration’ numbered list."""
    doc.add_heading('2. LLM Role and Integration', level=2)
    for _, row in df.iterrows():
        p = doc.add_paragraph(style='List Number')
        p.add_run(f"{row['paper_title']}: {row['llm_integration_type']}")

def add_language_section(doc: Document, df: pd.DataFrame):
    """Add the ‘State of the Field by Language’ section."""
    doc.add_heading('3. State of the Field by Language', level=2)
    for lang, group in df.groupby('programming_language'):
        bug_focus = ", ".join(sorted(group['primary_bug_type'].unique()))
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{lang}: {len(group)} papers, bug types: {bug_focus}")

def generate_summary_doc(excel_path: str, output_path: str):
    """Load data, build the Word doc, and save it."""
    df = load_data(excel_path)
    doc = Document()
    doc.add_heading('LLMs + Bug Detection Summary', level=1)

    add_general_section(doc, df)
    add_specific_section(doc, df)
    add_integration_section(doc, df)
    add_language_section(doc, df)

    doc.save(output_path)
    print(f"Written summary to {output_path}")

if __name__ == "__main__":
    EXCEL_FILE = 'papers.xlsx'
    OUTPUT_DOC  = 'bug_detection_summary.docx'
    generate_summary_doc(EXCEL_FILE, OUTPUT_DOC)